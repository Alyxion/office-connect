"""Unit tests for the document peek helpers."""

from __future__ import annotations

import io

import pytest

from office_con.peek import (
    classify,
    peek_docx,
    peek_document,
    peek_pdf,
    peek_xlsx,
)


class TestClassify:

    def test_pdf_by_extension(self):
        assert classify("report.pdf", None) == "pdf"

    def test_xlsx_by_extension(self):
        assert classify("sheet.xlsx", None) == "xlsx"
        assert classify("macro.xlsm", None) == "xlsx"

    def test_docx_by_extension(self):
        assert classify("draft.docx", None) == "docx"

    def test_pdf_by_mime(self):
        assert classify(None, "application/pdf") == "pdf"

    def test_xlsx_by_mime(self):
        assert classify(None,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet") == "xlsx"

    def test_docx_by_mime(self):
        assert classify(None,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document") == "docx"

    def test_unknown_returns_none(self):
        assert classify("foo.bar", "application/binary") is None
        assert classify(None, None) is None

    def test_case_insensitive(self):
        assert classify("REPORT.PDF", None) == "pdf"


class TestPeekPdf:

    def _tiny_pdf(self) -> bytes:
        """Make a tiny 1-page PDF on the fly using pypdfium2 helpers are
        complex; use reportlab-free approach: craft a minimal valid PDF."""
        # pdfplumber / pypdfium2 both accept this basic valid PDF.
        return (
            b"%PDF-1.4\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n"
            b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 200 200]/Contents 4 0 R/Resources<<>>>>endobj\n"
            b"4 0 obj<</Length 44>>stream\n"
            b"BT /F1 12 Tf 10 100 Td (Hello Peek) Tj ET\n"
            b"endstream endobj\n"
            b"xref\n0 5\n"
            b"0000000000 65535 f\n"
            b"0000000009 00000 n\n"
            b"0000000052 00000 n\n"
            b"0000000097 00000 n\n"
            b"0000000170 00000 n\n"
            b"trailer<</Size 5/Root 1 0 R>>\n"
            b"startxref\n260\n%%EOF\n"
        )

    def test_page_count_and_kind(self):
        result = peek_pdf(self._tiny_pdf(), render=False)
        assert result["kind"] == "pdf"
        assert result["page_count"] == 1
        assert len(result["pages"]) == 1

    def test_render_produces_png(self):
        result = peek_pdf(self._tiny_pdf(), render=True, render_scale=1.0)
        assert "renders" in result
        assert len(result["renders"]) == 1
        r = result["renders"][0]
        assert r["png_base64"]
        assert r["width"] > 0
        assert r["height"] > 0


class TestPeekXlsx:

    def _tiny_xlsx(self) -> bytes:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Main"
        ws.append(["Name", "Value"])
        ws.append(["alpha", 1])
        ws.append(["beta", 2])
        wb.create_sheet("Extra")
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_lists_sheets(self):
        result = peek_xlsx(self._tiny_xlsx())
        assert result["kind"] == "xlsx"
        assert "Main" in result["sheet_names"]
        assert "Extra" in result["sheet_names"]
        assert result["active_sheet"] == "Main"

    def test_samples_rows(self):
        result = peek_xlsx(self._tiny_xlsx())
        sheet = result["sheets"][0]
        assert sheet["name"] == "Main"
        assert sheet["rows"][0] == ["Name", "Value"]
        assert sheet["rows_shown"] == 3

    def test_row_cap(self):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        for i in range(100):
            ws.append([i])
        buf = io.BytesIO()
        wb.save(buf)
        result = peek_xlsx(buf.getvalue(), max_rows=10)
        assert result["sheets"][0]["rows_shown"] == 10

    def test_all_sheets(self):
        result = peek_xlsx(self._tiny_xlsx(), include_all_sheets=True)
        assert len(result["sheets"]) == 2


class TestPeekDocx:

    def _tiny_docx(self) -> bytes:
        import docx
        doc = docx.Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("This is the intro.")
        doc.add_heading("Section A", level=2)
        doc.add_paragraph("Section A body.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_basic(self):
        result = peek_docx(self._tiny_docx())
        assert result["kind"] == "docx"
        assert result["paragraph_count"] >= 4
        assert len(result["headings"]) == 2
        assert result["headings"][0]["text"] == "Title"

    def test_paragraph_cap(self):
        import docx
        doc = docx.Document()
        for i in range(100):
            doc.add_paragraph(f"para {i}")
        buf = io.BytesIO()
        doc.save(buf)
        result = peek_docx(buf.getvalue(), max_paragraphs=5)
        assert len(result["paragraphs"]) == 5


class TestPeekDocumentDispatch:

    def test_unsupported_reports_size(self):
        data = b"\x00" * 1234
        result = peek_document(data, name="weird.bin", content_type="application/octet-stream")
        assert result["kind"] == "unsupported"
        assert result["byte_size"] == 1234

    def test_dispatches_to_xlsx(self):
        import openpyxl
        wb = openpyxl.Workbook()
        wb.active.append(["a", "b"])
        buf = io.BytesIO()
        wb.save(buf)
        result = peek_document(buf.getvalue(), name="test.xlsx")
        assert result["kind"] == "xlsx"
