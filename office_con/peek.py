"""Document peek helpers — extract a compact summary from PDFs, xlsx, docx.

The goal is to let the MCP server tell the LLM what's IN an attachment
without shipping the full binary back through the tool result. Each peek
function takes bytes + options and returns a small dict (plus optional
PNG bytes for rendered pages).

Licenses of underlying libs (all non-AGPL):
    pypdfium2  Apache-2.0 / BSD-3
    pdfplumber MIT
    openpyxl   MIT
    python-docx MIT
"""

from __future__ import annotations

import base64
import io
from typing import Any


SUPPORTED_EXTENSIONS = {
    "pdf": "pdf",
    "xlsx": "xlsx",
    "xlsm": "xlsx",
    "docx": "docx",
}


def classify(name: str | None, content_type: str | None) -> str | None:
    """Return our internal kind ('pdf'/'xlsx'/'docx') or None if unsupported.

    Prefers the filename extension; falls back to the MIME type.
    """
    if name:
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        if ext in SUPPORTED_EXTENSIONS:
            return SUPPORTED_EXTENSIONS[ext]
    if content_type:
        ct = content_type.lower()
        if "pdf" in ct:
            return "pdf"
        if "spreadsheetml" in ct or "excel" in ct:
            return "xlsx"
        if "wordprocessingml" in ct or "msword" in ct:
            return "docx"
    return None


def _truncate(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + f"\n… [truncated: {len(text) - limit} more chars]"


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def peek_pdf(
    data: bytes,
    *,
    pages: int = 1,
    render: bool = True,
    render_scale: float = 1.5,
    text_limit_per_page: int = 4000,
) -> dict[str, Any]:
    """Peek at a PDF.

    Returns a dict with:
      - page_count, metadata
      - pages: list of {index, text}
      - renders: list of {index, png_base64, width, height} if render=True
    """
    import pdfplumber
    import pypdfium2 as pdfium

    result: dict[str, Any] = {"kind": "pdf"}

    # Metadata + text via pdfplumber (reading order is better than pypdfium2).
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        result["page_count"] = len(pdf.pages)
        result["metadata"] = dict(pdf.metadata or {})
        n = min(pages, len(pdf.pages))
        extracted_pages = []
        for i in range(n):
            try:
                text = pdf.pages[i].extract_text() or ""
            except Exception as exc:
                text = f"[pdfplumber extract failed: {exc}]"
            extracted_pages.append({
                "index": i,
                "text": _truncate(text, text_limit_per_page),
            })
        result["pages"] = extracted_pages

    # Renders via pypdfium2 (fast, uses Chrome's PDF engine).
    if render:
        renders = []
        doc = pdfium.PdfDocument(data)
        try:
            n_render = min(pages, len(doc))
            for i in range(n_render):
                page = doc[i]
                try:
                    pil = page.render(scale=render_scale).to_pil()
                    buf = io.BytesIO()
                    pil.save(buf, format="PNG", optimize=True)
                    png_bytes = buf.getvalue()
                    renders.append({
                        "index": i,
                        "png_base64": base64.b64encode(png_bytes).decode("ascii"),
                        "width": pil.width,
                        "height": pil.height,
                        "byte_size": len(png_bytes),
                    })
                finally:
                    page.close()
        finally:
            doc.close()
        result["renders"] = renders

    return result


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def peek_xlsx(
    data: bytes,
    *,
    max_rows: int = 30,
    max_cols: int = 30,
    include_all_sheets: bool = False,
) -> dict[str, Any]:
    """Peek at an xlsx workbook.

    Returns sheet names, metadata, and a sample of cells from the active
    sheet (or all sheets if ``include_all_sheets=True``).
    """
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sheets = []
    names = wb.sheetnames
    target_sheets = names if include_all_sheets else [wb.active.title] if wb.active else []

    for sheet_name in target_sheets:
        ws = wb[sheet_name]
        rows: list[list[Any]] = []
        n_rows = 0
        for row in ws.iter_rows(values_only=True):
            n_rows += 1
            if n_rows > max_rows:
                break
            rows.append(list(row[:max_cols]))
        sheets.append({
            "name": sheet_name,
            "rows_shown": len(rows),
            "max_col": ws.max_column,
            "max_row": ws.max_row,
            "rows": rows,
        })

    result = {
        "kind": "xlsx",
        "sheet_names": names,
        "active_sheet": wb.active.title if wb.active else None,
        "sheets": sheets,
    }
    try:
        props = wb.properties
        result["metadata"] = {
            "title": props.title,
            "creator": props.creator,
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None,
        }
    except Exception:
        pass
    wb.close()
    return result


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def peek_docx(
    data: bytes,
    *,
    max_paragraphs: int = 30,
    text_limit_per_paragraph: int = 400,
) -> dict[str, Any]:
    """Peek at a docx. Returns heading/paragraph samples and metadata."""
    import docx

    doc = docx.Document(io.BytesIO(data))
    paragraphs = []
    headings = []
    for i, p in enumerate(doc.paragraphs):
        if len(paragraphs) >= max_paragraphs:
            break
        text = (p.text or "").strip()
        if not text:
            continue
        style = p.style.name if p.style else None
        entry = {
            "index": i,
            "style": style,
            "text": _truncate(text, text_limit_per_paragraph),
        }
        paragraphs.append(entry)
        if style and style.lower().startswith("heading"):
            headings.append({"index": i, "style": style, "text": text})

    result: dict[str, Any] = {
        "kind": "docx",
        "paragraph_count": sum(1 for p in doc.paragraphs if (p.text or "").strip()),
        "total_paragraphs": len(doc.paragraphs),
        "headings": headings,
        "paragraphs": paragraphs,
        "table_count": len(doc.tables),
    }
    try:
        props = doc.core_properties
        result["metadata"] = {
            "title": props.title,
            "author": props.author,
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None,
        }
    except Exception:
        pass
    return result


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------


def peek_document(
    data: bytes,
    *,
    name: str | None = None,
    content_type: str | None = None,
    pages: int = 1,
    render: bool = True,
    max_rows: int = 30,
    max_paragraphs: int = 30,
    all_sheets: bool = False,
) -> dict[str, Any]:
    """Dispatch a peek based on filename extension / MIME. Returns a dict
    including a ``kind`` field ('pdf'|'xlsx'|'docx'|'unsupported')."""
    kind = classify(name, content_type)
    if kind is None:
        return {
            "kind": "unsupported",
            "name": name,
            "content_type": content_type,
            "byte_size": len(data),
            "message": (
                "File type not supported for peek. Supported: pdf, xlsx, docx. "
                "Use o365_get_file_content or fetch the attachment directly to "
                "retrieve raw bytes."
            ),
        }
    if kind == "pdf":
        return peek_pdf(data, pages=pages, render=render)
    if kind == "xlsx":
        return peek_xlsx(data, max_rows=max_rows, include_all_sheets=all_sheets)
    if kind == "docx":
        return peek_docx(data, max_paragraphs=max_paragraphs)
    return {"kind": "unsupported", "name": name, "content_type": content_type}
